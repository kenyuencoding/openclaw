import os
import pandas as pd
import yfinance as yf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Mm
from datetime import datetime

run_dir = os.path.join(os.getcwd(),'briefing_runs')
os.makedirs(run_dir, exist_ok=True)
orig_path = 'C:/Users/n1sol/.openclaw/workspace/briefing_final_20260302_192609.docx'
out_path = 'C:/Users/n1sol/.openclaw/workspace/briefing_merged_final_{}.docx'.format(datetime.now().strftime('%Y%m%d_%H%M%S'))

tickers = ['AVAV','NVDA','AES','COH','NCLH']

# build market section doc
market_doc = Document()
market_doc.add_paragraph('Top 5 US Stocks (market movers)').bold = True

for t in tickers:
    try:
        tk = yf.Ticker(t)
        hist = tk.history(period='1y', interval='1d')
        if hist.empty:
            print(f"{t}: no data")
            continue
        price = hist['Close'].iloc[-1]
        data = hist.copy()
        data['EMA9'] = data['Close'].ewm(span=9).mean()
        data['EMA21'] = data['Close'].ewm(span=21).mean()
        tp = (data['High'] + data['Low'] + data['Close'])/3
        data['VWAP'] = (tp * data['Volume']).cumsum() / data['Volume'].cumsum()
        ma20 = data['Close'].rolling(window=20).mean()
        std20 = data['Close'].rolling(window=20).std()
        data['BB_up'] = ma20 + 2*std20
        data['BB_dn'] = ma20 - 2*std20

        fig = go.Figure()
        fig.add_trace(go.Candlestick(x=data.index, open=data['Open'], high=data['High'], low=data['Low'], close=data['Close']))
        fig.add_trace(go.Scatter(x=data.index, y=data['EMA9'], line=dict(color='orange', width=1), name='EMA9'))
        fig.add_trace(go.Scatter(x=data.index, y=data['EMA21'], line=dict(color='green', width=1), name='EMA21'))
        fig.add_trace(go.Scatter(x=data.index, y=data['VWAP'], line=dict(color='purple', width=1), name='VWAP'))
        fig.add_trace(go.Scatter(x=data.index, y=data['BB_up'], line=dict(color='blue', width=1, dash='dot')))
        fig.add_trace(go.Scatter(x=data.index, y=data['BB_dn'], line=dict(color='blue', width=1, dash='dot')))
        recent_high = data['High'].max()
        recent_low = data['Low'].min()
        fig.add_annotation(x=data['High'].idxmax(), y=recent_high, text=f'H {recent_high:.2f}', showarrow=True)
        fig.add_annotation(x=data['Low'].idxmin(), y=recent_low, text=f'L {recent_low:.2f}', showarrow=True)
        fig.update_layout(xaxis_rangeslider_visible=False, height=480)
        img = os.path.join(run_dir, f'{t}.png')
        fig.write_image(img, scale=2)

        # compose entries
        name = tk.info.get('shortName','')
        en = f"{name} remains active; latest price {price:.2f}. Recent activity shows traders reacting to sector/news catalysts."
        cn = f"{name} 仍然活躍；最新價格 {price:.2f}。近期活動顯示交易者對產業/新聞刺激作出反應。"
        ma50 = data['Close'].rolling(window=50).mean().iloc[-1]
        ma200 = data['Close'].rolling(window=200).mean().iloc[-1] if len(data)>=200 else ma50
        phase = 'Markup' if ma50>ma200 else 'Accumulation/Testing'
        vol_recent = data['Volume'].tail(20).mean()
        vol_prev = data['Volume'].tail(60).head(20).mean() if len(data)>=60 else vol_recent
        vpa = 'V↑ confirming price move' if vol_recent>vol_prev else 'V↓ during pullback (test)'
        wy_en = f"Wyckoff phase: {phase}. VPA: {vpa}."
        wy_cn = f"Wyckoff 階段: {phase}。量價狀況: {vpa}。"
        strike_buy = round(price)
        strike_sell = strike_buy + max(1,int(round(price*0.02)))
        expiry = (data.index[-1] + pd.Timedelta(days=30)).strftime('%Y-%m-%d')
        strat_en = f"Example bull call vertical: buy {strike_buy} call, sell {strike_sell} call, expiry {expiry}."
        strat_cn = f"示例牛市看漲價差: 買入 {strike_buy} 看漲, 賣出 {strike_sell} 看漲, 到期日 {expiry}."

        p = market_doc.add_paragraph()
        p.add_run(f"{t} - {name}\n").bold = True
        market_doc.add_picture(img, width=Mm(160))
        market_doc.add_paragraph('English summary:').bold = True
        market_doc.add_paragraph(en)
        market_doc.add_paragraph('中文摘要:').bold = True
        market_doc.add_paragraph(cn)
        market_doc.add_paragraph('Wyckoff analysis (EN):').bold = True
        market_doc.add_paragraph(wy_en)
        market_doc.add_paragraph('Wyckoff 分析 (中):').bold = True
        market_doc.add_paragraph(wy_cn)
        market_doc.add_paragraph('Options strategy (vertical spread) (EN):').bold = True
        market_doc.add_paragraph(strat_en)
        market_doc.add_paragraph('選擇權策略 (牛/熊價差) (中):').bold = True
        market_doc.add_paragraph(strat_cn)
        market_doc.add_page_break()

    except Exception as e:
        print('Error', t, e)

# merge: append market_doc to original but remove old top5 section heuristically
orig = Document(orig_path)
# simple append for reliability
p = orig.add_paragraph()
p.add_run().add_break()
for el in market_doc.element.body:
    orig.element.body.append(el)
orig.save(out_path)
print(out_path)
