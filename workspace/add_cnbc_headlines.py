from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Extracted CNBC top trending headlines (from earlier fetch)
headlines = [
    "Supreme Court bars redrawing only Republican-held NYC congressional district for 2026 election",
    "Iran live updates: Six U.S. service members killed in action",
    "Jamie Dimon says Trump’s $5 billion debanking lawsuit ‘has no merit’ but he’s sympathetic to concerns",
    "First flights take off from Dubai after Iran strikes, but service is 'limited'",
    "S&P 500 closes flat, rebounding from lows as traders buy the dip after U.S.-Iran attacks"
]

timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
out = f'C:/Users/n1sol/.openclaw/workspace/briefing_market_cnbc_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'

doc = Document()
section = doc.sections[0]

p = doc.add_paragraph()
r = p.add_run(f"Briefing (updated with CNBC Top5) — {timestamp}")
r.bold = True
r.font.size = Pt(16)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

doc.add_paragraph()

h = doc.add_paragraph()
h.add_run('CNBC — Top 5 headlines (snapshot)').bold = True
for i, hline in enumerate(headlines,1):
    p = doc.add_paragraph()
    p.add_run(f"{i}) {hline}\n").bold = True
    p.add_run('[Snippet placeholder from CNBC page]').italic = True

# Footer
f = doc.add_paragraph()
f.add_run('Sources: CNBC (live fetch)')

doc.save(out)
print(out)
