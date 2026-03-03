from docx import Document
from datetime import datetime

base='C:/Users/n1sol/.openclaw/workspace/briefing_final_20260302_192609.docx'
add='C:/Users/n1sol/.openclaw/workspace/briefing_market_cnbc_filled_20260303_080018.docx'
out='C:/Users/n1sol/.openclaw/workspace/briefing_merged_20260302_{}.docx'.format(datetime.now().strftime('%Y%m%d_%H%M%S'))

doc1=Document(base)
doc2=Document(add)

# append page break
p = doc1.add_paragraph()
run = p.add_run()
run.add_break()

for element in doc2.element.body:
    doc1.element.body.append(element)

doc1.save(out)
print(out)
