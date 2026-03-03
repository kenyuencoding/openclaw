import os
from datetime import datetime
import yfinance as yf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

RUN_DIR = os.path.join(os.getcwd(), 'briefing_runs')
os.makedirs(RUN_DIR, exist_ok=True)

timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
out_docx = os.path.join(os.getcwd(), f'briefing_market_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

# Example tickers - replace with dynamic selection as needed
tickers = ['NVDA','NFLX','NU','PLUG','SOFI']

doc = Document()
section = doc.sections[0]
section.left_margin = Mm(12)
section.right_margin = Mm(12)

p = doc.add_paragraph()
r = p.add_run(f"Market Movers Briefing — {timestamp}")
r.bold = True
r.font.size = Pt(18)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_ticker_section(ticker):
    tk = yf.Ticker(ticker)
    hist = tk.history(period='60d', interval='1d')
    if hist.empty:
        return
    # Use last 30 days for chart clarity
    data = hist.tail(30)
    fig = go.Figure()
    fig.add_trace(go.Candlestick(x=data.index,
                                 open=data['Open'], high=data['High'], low=data['Low'], close=data['Close'], name='Price'))
    # EMA9/21
    data['EMA9'] = data['Close'].ewm(span=9).mean()
    data['EMA21'] = data['Close'].ewm(span=21).mean()
    fig.add_trace(go.Scatter(x=data.index, y=data['EMA9'], line=dict(color='orange', width=1), name='EMA9'))
    fig.add_trace(go.Scatter(x=data.index, y=data['EMA21'], line=dict(color='green', width=1), name='EMA21'))
    # VWAP approximate using typical price * volume cum / vol cum
    tp = (data['High'] + data['Low'] + data['Close'])/3
    vwap = (tp * data['Volume']).cumsum() / data['Volume'].cumsum()
    fig.add_trace(go.Scatter(x=data.index, y=vwap, line=dict(color='purple', width=1), name='VWAP'))
    # volume
    fig.update_layout(xaxis_rangeslider_visible=False, height=360)
    # save image
    img_path = os.path.join(RUN_DIR, f'{ticker}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.png')
    fig.write_image(img_path, scale=2)

    # Add to doc
    h = doc.add_paragraph()
    hr = h.add_run(f"{ticker} - {tk.info.get('shortName', '')}\n")
    hr.bold = True
    hr.font.size = Pt(13)
    doc.add_picture(img_path, width=Mm(160))
    # Add placeholder analysis
    p = doc.add_paragraph()
    p.add_run("English summary:\n").bold = True
    p.add_run("[Auto summary here]\n")
    p.add_run("中文摘要:\n").bold = True
    p.add_run("[中文摘要在此]\n")
    p.add_run("Wyckoff analysis (EN/中):\n").bold = True
    p.add_run("[Wyckoff analysis here]\n")
    p.add_run("Options strategy (EN/中):\n").bold = True
    p.add_run("[Suggested vertical spread here]\n")

for t in tickers:
    try:
        add_ticker_section(t)
    except Exception as e:
        print('Error for', t, e)

# Footer
f = doc.add_paragraph()
f.add_run('Sources: Yahoo Finance (prices), compiled by Yuetki')

doc.save(out_docx)
print(out_docx)
