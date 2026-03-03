import os
from datetime import datetime
import re
import requests
import yfinance as yf
import plotly.graph_objects as go
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

RUN_DIR = os.path.join(os.getcwd(), 'briefing_runs')
os.makedirs(RUN_DIR, exist_ok=True)

# Fetch CNBC article that lists tickers
url = 'https://www.cnbc.com/2026/03/02/stocks-making-the-biggest-moves-midday-avav-nvda-aes-coh-nclh.html'
try:
    r = requests.get(url, timeout=15)
    text = r.text
except Exception as e:
    text = ''

# Extract tickers from the URL or text as fallback
match = re.search(r'-([a-z0-9-]+)\.html$', url)
if match:
    slug = match.group(1)
    parts = slug.split('-')
    # last segment contains tickers separated by '-'
    tickers = [p.upper() for p in parts if p.isalpha() or p.isalnum()]
else:
    tickers = ['NVDA','NFLX','NU','PLUG','SOFI']

# Clean tickers by removing words
clean = []
for p in parts:
    if len(p) <= 5:
        clean.append(p.upper())
if clean:
    tickers = clean[:5]
else:
    tickers = ['NVDA','NFLX','NU','PLUG','SOFI']

# Build DOCX
timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
out_docx = os.path.join(os.getcwd(), f'briefing_market_cnbc_full_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

doc = Document()
section = doc.sections[0]
section.left_margin = Mm(12)
section.right_margin = Mm(12)

p = doc.add_paragraph()
r = p.add_run(f"Market Movers Briefing (CNBC top list) — {timestamp}")
r.bold = True
r.font.size = Pt(18)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

for t in tickers:
    try:
        tk = yf.Ticker(t)
        hist = tk.history(period='1y', interval='1d')
        if hist.empty:
            continue
        data = hist.tail(365)
        # compute indicators
        data['EMA9'] = data['Close'].ewm(span=9).mean()
        data['EMA21'] = data['Close'].ewm(span=21).mean()
        tp = (data['High'] + data['Low'] + data['Close'])/3
        vwap = (tp * data['Volume']).cumsum() / data['Volume'].cumsum()
        data['VWAP'] = vwap
        # Bollinger bands
        ma20 = data['Close'].rolling(window=20).mean()
        std20 = data['Close'].rolling(window=20).std()
        data['BB_up'] = ma20 + 2*std20
        data['BB_dn'] = ma20 - 2*std20

        fig = go.Figure()
        fig.add_trace(go.Candlestick(x=data.index, open=data['Open'], high=data['High'], low=data['Low'], close=data['Close'], name='Price'))
        fig.add_trace(go.Scatter(x=data.index, y=data['EMA9'], line=dict(color='orange', width=1), name='EMA9'))
        fig.add_trace(go.Scatter(x=data.index, y=data['EMA21'], line=dict(color='green', width=1), name='EMA21'))
        fig.add_trace(go.Scatter(x=data.index, y=data['VWAP'], line=dict(color='purple', width=1), name='VWAP'))
        fig.add_trace(go.Scatter(x=data.index, y=data['BB_up'], line=dict(color='rgba(0,0,255,0.2)'), name='BB_up'))
        fig.add_trace(go.Scatter(x=data.index, y=data['BB_dn'], line=dict(color='rgba(0,0,255,0.2)'), name='BB_dn'))
        # volume as bar
        fig.update_layout(xaxis_rangeslider_visible=False, height=480)
        img_path = os.path.join(RUN_DIR, f'{t}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.png')
        fig.write_image(img_path, scale=2)
        # add to doc
        h = doc.add_paragraph()
        hr = h.add_run(f"{t} - {tk.info.get('shortName','')}\n")
        hr.bold = True
        hr.font.size = Pt(13)
        doc.add_picture(img_path, width=Mm(160))
        # add placeholders
        p = doc.add_paragraph()
        p.add_run("English summary:\n").bold = True
        p.add_run("[Auto summary will be filled]\n")
        p.add_run("中文摘要:\n").bold = True
        p.add_run("[中文摘要將會填寫]\n")
        p.add_run("Wyckoff analysis (EN/中):\n").bold = True
        p.add_run("[Wyckoff analysis]\n")
        p.add_run("Options strategy (vertical spread) (EN/中):\n").bold = True
        p.add_run("[Suggested vertical spread]\n")
    except Exception as e:
        print('error', t, e)

# footer
f = doc.add_paragraph()
f.add_run('Sources: CNBC (top movers), Yahoo Finance (price data). Compiled by Yuetki')

doc.save(out_docx)
print(out_docx)
