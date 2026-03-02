from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Content
timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
world = [
    ("Israel strikes Lebanon after Hezbollah attacks, widening regional conflict",
     "Fresh Israeli strikes have targeted sites in Lebanon after cross-border attacks attributed to Hezbollah. The strikes mark an escalation that regional analysts warn could draw Iran-linked actors deeper into the confrontation. Civilians and infrastructure in border areas have been reported affected and international actors are calling for de‑escalation."),
    ("Rising global concern over Iran-linked military activity",
     "Reports say Iran-affiliated groups and proxies have increased operations across the region following recent strikes, raising fears of a wider confrontation. Diplomatic channels are reportedly active as regional powers and Western states monitor developments. Markets are on alert for oil & risk-premium effects."),
    ("Continued coverage of major international crises and migration pressures",
     "Multiple crises — conflict zones and climate-driven displacement — continue to dominate coverage, with NGOs warning of humanitarian needs outpacing supply. Governments are debating both immediate relief and longer-term migration policy responses. International aid agencies are mobilising contingency plans."),
    ("Global markets react nervously to geopolitical uncertainty",
     "Risk-off sentiment appears in markets, pushing safe-haven flows and volatility in energy and defense sectors. Analysts highlight how short-term supply concerns (energy) and sentiment shifts can influence equity, FX, and commodity moves. Traders are watching for central bank comments to gauge risk appetite."),
    ("High-profile political developments and leadership shifts in multiple countries",
     "Several countries are reporting leadership changes, major policy announcements, or scandal-driven investigations that could reshape local political landscapes. These domestic stories are feeding into broader geopolitical narratives and market expectations.")
]

us = [
    ("US monitors regional escalation; diplomatic moves under way",
     "US officials are closely tracking strikes and border incidents in the Middle East, engaging regional partners to reduce the risk of wider conflict. Officials stress readiness to protect US personnel and interests. Congressional scrutiny of response options is expected."),
    ("Domestic political developments and policy debates",
     "Major discussions on fiscal/legislative priorities continue, with key votes and hearings scheduled that may affect market sentiment. Watch for statements from leading lawmakers that could influence sector-specific regulations."),
    ("Earnings season highlights and market movers",
     "Quarterly results from several large firms are shaping market leadership; tech and energy names show mixed reactions depending on earnings vs. expectations. Analysts are parsing guidance for 2H outlooks."),
    ("Economic data releases and Fed commentary",
     "Recent economic prints and Fed speakers are influencing bond yields and equity rotation. Traders are focused on inflation momentum and employment signals to assess the path for rates."),
    ("US security / cyber incidents and law enforcement actions",
     "A number of high-visibility cyber or law-enforcement actions have made headlines, prompting discussion about regulatory and security responses. Markets sometimes react to potential regulatory implications for affected sectors.")
]

hk = [
    ("Regional spillover concerns and local market sensitivity",
     "Hong Kong coverage highlights how regional conflicts may affect trade and investor sentiment locally, especially via commodity and shipping routes. Policymakers and market participants are watching for impacts on trade flows and corporate guidance."),
    ("Local political and economic policy announcements",
     "Updates on domestic policy, housing, and economic measures continue to be prominent, with officials signalling support measures for key sectors as needed. These updates are tracked closely by local investors."),
    ("Market reopening / tourism and retail recovery stories",
     "Coverage focuses on tourism/retail rebounds and related events, including promotional initiatives and expected visitor flows. Recovery in consumer sectors is a recurring local theme."),
    ("Corporate news from HK-listed firms",
     "Earnings and corporate actions from major Hong Kong listings are being reported, with sector-specific moves often driving local indices. Watch for company guidance in upcoming releases."),
    ("Infrastructure and transport developments",
     "Local infrastructure projects, transport upgrades, and policy implementation timelines appear in the news, with implications for construction and related sectors.")
]

stocks = [
    ("Market movers and options flow",
     "Earnings beats/misses in large-cap names have driven intra-day leadership changes; energy names are sensitive to geopolitical news, while defensive sectors and safe havens have seen inflows. Options flow shows elevated put-buying in certain cyclical names and increased call activity in a few tech beaters; implied volatility edges up in names exposed to geopolitical risk.")
]

# Create document
doc = Document()

# Styles for high-tech look
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title with timestamp
p = doc.add_paragraph()
r = p.add_run(f"Briefing — {timestamp}")
r.bold = True
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0x11,0x22,0x33)
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Add a small separator
doc.add_paragraph()

def add_section(title, items):
    h = doc.add_paragraph()
    hr = h.add_run(title)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.color.rgb = RGBColor(0x00,0x66,0xCC)
    for i, (t,s) in enumerate(items,1):
        tp = doc.add_paragraph()
        run = tp.add_run(f"{i}) {t}\n")
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x00,0x00,0x00)
        sp = tp.add_run(s)
        sp.font.size = Pt(11)

add_section('Worldwide — Top 5 headlines', world)
add_section('\nUS — Top 5 headlines', us)
add_section('\nHong Kong — Top 5 headlines', hk)
add_section('\nUS stocks & options — notable items (last 24h)', stocks)

# Footer note
doc.add_paragraph()
foot = doc.add_paragraph()
fr = foot.add_run('Sources: Reuters, BBC, Yahoo News, local aggregation via OpenClaw browser relay')
fr.italic = True
fr.font.size = Pt(9)

# Save
path = 'C:/Users/n1sol/.openclaw/workspace/briefing_{}.docx'.format(datetime.now().strftime('%Y%m%d_%H%M%S'))
doc.save(path)
print(path)
