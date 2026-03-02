from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Mm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')

# English + Chinese content (manual translations)
world = [
    ("Israel strikes Lebanon after Hezbollah attacks, widening regional conflict",
     "Fresh Israeli strikes have targeted sites in Lebanon after cross-border attacks attributed to Hezbollah. The strikes mark an escalation that regional analysts warn could draw Iran-linked actors deeper into the confrontation. Civilians and infrastructure in border areas have been reported affected and international actors are calling for de‑escalation.",
     "以色列在真主黨襲擊後對黎巴嫩進行空襲，地區衝突擴大。以色列最新空襲鎖定黎巴嫩目標，分析人士警告此舉可能將與伊朗有關的行動者更深地捲入衝突。邊境地區已有平民與基礎設施受影響，國際社會呼籲降級局勢。"),
    ("Rising global concern over Iran-linked military activity",
     "Reports say Iran-affiliated groups and proxies have increased operations across the region following recent strikes, raising fears of a wider confrontation. Diplomatic channels are reportedly active as regional powers and Western states monitor developments. Markets are on alert for oil & risk-premium effects.",
     "報導稱與伊朗有關的團體與代理人於近期攻擊後在區域內加強行動，外界擔憂衝突擴大。地區強權與西方國家透過外交渠道密切關注事態。市場對油價與風險溢價敏感。"),
    ("Continued coverage of major international crises and migration pressures",
     "Multiple crises — conflict zones and climate-driven displacement — continue to dominate coverage, with NGOs warning of humanitarian needs outpacing supply. Governments are debating both immediate relief and longer-term migration policy responses. International aid agencies are mobilising contingency plans.",
     "多項危機——包括衝突地區與氣候引發的人口流離失所——持續成為報導焦點，非政府組織警告人道需求超過供給。各國政府在討論短期救援與長期移民政策應對。國際救援機構啟動應急計畫。"),
    ("Global markets react nervously to geopolitical uncertainty",
     "Risk-off sentiment appears in markets, pushing safe-haven flows and volatility in energy and defense sectors. Analysts highlight how short-term supply concerns (energy) and sentiment shifts can influence equity, FX, and commodity moves. Traders are watching for central bank comments to gauge risk appetite.",
     "由於地緣政治不確定性，市場出現避險情緒，資金流入安全資產，使能源與國防類股波動加劇。分析師指出短期供應擔憂與情緒變動會影響股票、外匯與大宗商品走勢。交易員關注央行言論以判斷風險偏好。"),
    ("High-profile political developments and leadership shifts in multiple countries",
     "Several countries are reporting leadership changes, major policy announcements, or scandal-driven investigations that could reshape local political landscapes. These domestic stories are feeding into broader geopolitical narratives and market expectations.",
     "多國報導領導層變動、重要政策宣布或醜聞調查，可能改寫當地政治版圖。這些國內新聞也成為廣泛地緣政治敘事與市場預期的一部分。")
]

us = [
    ("US monitors regional escalation; diplomatic moves under way",
     "US officials are closely tracking strikes and border incidents in the Middle East, engaging regional partners to reduce the risk of wider conflict. Officials stress readiness to protect US personnel and interests. Congressional scrutiny of response options is expected.",
     "美國密切監視地區升級情勢，並與區域夥伴互動以降低衝突擴大的風險。官員強調保護美國人員與利益的準備。國會可能對因應方案進行審查。"),
    ("Domestic political developments and policy debates",
     "Major discussions on fiscal/legislative priorities continue, with key votes and hearings scheduled that may affect market sentiment. Watch for statements from leading lawmakers that could influence sector-specific regulations.",
     "國內重大財政與立法議題持續討論，關鍵表決與聽證會可能影響市場情緒。注意重要立法者的聲明，可能牽動特定產業法規。"),
    ("Earnings season highlights and market movers",
     "Quarterly results from several large firms are shaping market leadership; tech and energy names show mixed reactions depending on earnings vs. expectations. Analysts are parsing guidance for 2H outlooks.",
     "多家大型企業季報影響市場領導股表現；科技與能源類股依獲利與預期出現不同反應。分析師正在解讀下半年的展望指引。"),
    ("Economic data releases and Fed commentary",
     "Recent economic prints and Fed speakers are influencing bond yields and equity rotation. Traders are focused on inflation momentum and employment signals to assess the path for rates.",
     "近期數據與聯準會官員談話影響債券殖利率與股票輪動。交易員側重通膨趨勢與就業訊號以評估利率走向。"),
    ("US security / cyber incidents and law enforcement actions",
     "A number of high-visibility cyber or law-enforcement actions have made headlines, prompting discussion about regulatory and security responses. Markets sometimes react to potential regulatory implications for affected sectors.",
     "數起高調的網安或執法行動成為頭條，促使對法規與安全回應的討論。市場有時會對受影響產業可能的法規影響做出反應。")
]

hk = [
    ("Regional spillover concerns and local market sensitivity",
     "Hong Kong coverage highlights how regional conflicts may affect trade and investor sentiment locally, especially via commodity and shipping routes. Policymakers and market participants are watching for impacts on trade flows and corporate guidance.",
     "香港報導關注區域衝突如何透過商品與航運路線影響當地貿易與投資人情緒。政策制定者與市場參與者密切觀察對貿易流與公司指引的影響。"),
    ("Local political and economic policy announcements",
     "Updates on domestic policy, housing, and economic measures continue to be prominent, with officials signalling support measures for key sectors as needed. These updates are tracked closely by local investors.",
     "國內政策、房屋與經濟措施的更新持續受到關注，官員在必要時表示將提供重點產業支持。當地投資人密切跟進。"),
    ("Market reopening / tourism and retail recovery stories",
     "Coverage focuses on tourism/retail rebounds and related events, including promotional initiatives and expected visitor flows. Recovery in consumer sectors is a recurring local theme.",
     "報導聚焦旅遊與零售回溫，包括宣傳活動與預期訪客量。消費類產業的復甦是持續關注的主題。"),
    ("Corporate news from HK-listed firms",
     "Earnings and corporate actions from major Hong Kong listings are being reported, with sector-specific moves often driving local indices. Watch for company guidance in upcoming releases.",
     "本地上市公司之財報與公司行動持續報導，產業動向常左右本地指數。留意公司接下來的指引。"),
    ("Infrastructure and transport developments",
     "Local infrastructure projects, transport upgrades, and policy implementation timelines appear in the news, with implications for construction and related sectors.",
     "本地基礎建設工程、運輸升級與政策實施時程進入新聞，對建築與相關產業具影響。")
]

stocks = [
    ("Market movers and options flow",
     "Earnings beats/misses in large-cap names have driven intra-day leadership changes; energy names are sensitive to geopolitical news, while defensive sectors and safe havens have seen inflows. Options flow shows elevated put-buying in certain cyclical names and increased call activity in a few tech beaters; implied volatility edges up in names exposed to geopolitical risk.",
     "大型股的盈餘優於或不及預期導致盤中領導股變動；能源股對地緣政治消息敏感，防禦性產業與避險資產則出現資金流入。選擇權交易顯示部分週期股買入看跌權增加，少數科技強勢股則有買入看漲權；受地緣風險影響的個股隱含波動率上升。")
]

# Create bilingual document with larger fonts and modern style

doc = Document()
section = doc.sections[0]
section.page_height = Mm(297)
section.page_width = Mm(210)
section.left_margin = Mm(12)
section.right_margin = Mm(12)

# Title
p = doc.add_paragraph()
r = p.add_run(f"Briefing — {timestamp}")
r.bold = True
r.font.size = Pt(18)
r.font.name = 'Arial'
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Subtitle (Chinese)
p2 = doc.add_paragraph()
r2 = p2.add_run(f"簡報 — {timestamp}")
r2.italic = True
r2.font.size = Pt(12)
r2.font.name = 'Microsoft JhengHei'
p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

def add_bilingual_section(title_en, items):
    h = doc.add_paragraph()
    hr = h.add_run(title_en)
    hr.bold = True
    hr.font.size = Pt(14)
    hr.font.name = 'Segoe UI'
    hr.font.color.rgb = RGBColor(0x00,0x88,0xFF)
    for i, (t_en, s_en, s_cn) in enumerate(items,1):
        tp = doc.add_paragraph()
        run = tp.add_run(f"{i}) {t_en}\n")
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = 'Segoe UI'
        sp = tp.add_run(s_en + "\n")
        sp.font.size = Pt(12)
        sp.font.name = 'Calibri'
        sp_cn = tp.add_run(s_cn)
        sp_cn.font.size = Pt(12)
        sp_cn.font.name = 'Microsoft JhengHei'
        sp_cn.font.color.rgb = RGBColor(0x33,0x33,0x33)

add_bilingual_section('Worldwide — Top 5 headlines', world)
add_bilingual_section('US — Top 5 headlines', us)
add_bilingual_section('Hong Kong — Top 5 headlines', hk)
add_bilingual_section('US stocks & options — notable items (last 24h)', stocks)

# Footer
f = doc.add_paragraph()
f_run = f.add_run('Sources: Reuters, BBC, Yahoo News, local aggregation via OpenClaw browser relay')
f_run.italic = True
f_run.font.size = Pt(10)

out = f'C:/Users/n1sol/.openclaw/workspace/briefing_bilingual_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
doc.save(out)
print(out)
